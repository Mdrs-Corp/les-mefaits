from docx import Document
from docx.shared import RGBColor,Pt
import random
document = Document()


filemig=open('alexis.bmp','rb')
mig=filemig.read()
filemig.close()
asci=" .:-=+*#%@"[::-1]
w=int.from_bytes(mig[18:22],byteorder='little')
mig=mig[54:]
p = document.add_paragraph('')
p.paragraph_format.line_spacing = 0.5
def tobyte(nb:int):
    he=hex(nb)[2:]
    while len(he)<2:
        he="0"+he
    return bytes.fromhex(he)


for i in range(len(mig)-3,0,-3):
    truc=(mig[i]+mig[i+1]+mig[i+2])//3
    run=p.add_run(asci[int(truc/255*(len(asci)-1))])
    font = run.font
    font.color.rgb = RGBColor(mig[i+2],mig[i+1],mig[i])
    font.size = Pt(5)
    font.name="Consolas"
    if i%w==0:
        p.add_run('\n')

title=''
for i in range(10):title+=random.choice('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ')
document.save(f'./rendu/{title}.docx')
print(title)
